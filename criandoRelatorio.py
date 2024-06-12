from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
agora = datetime.now()
hora_atual = agora.strftime("_%Y-%m-%d_%H-%M-%S")

def criar_relatorio(Tituloo,periodo,Nome,DataDoAtendimento):
    # Criando um novo documento
    documento = Document()

    # Título do relatório
    titulo = documento.add_heading(f'{Tituloo}', level=1)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtítulo
    subtitulo = documento.add_heading(f'{periodo}', level=2)
    subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introdução
    documento.add_heading('1. Introdução', level=2)
    documento.add_paragraph(
        'Este relatório apresenta os resultados de vendas do período de Janeiro 2024 a Março 2024.'
    )

    # Tabela de dados
    documento.add_heading('2. Dados Extraidos dos Cards', level=2)
    tabela = documento.add_table(rows=1, cols=3)
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Nome do MEI '
    hdr_cells[1].text = 'Data do Atendimento '
    hdr_cells[2].text = 'Receita'

    dados_vendas = [
        ('Janeiro', '1000', 'R$ 50.000'),
        ('Fevereiro', '1200', 'R$ 60.000'),
        ('Março', '1100', 'R$ 55.000')
    ]
    
    for i in range(0,len(Nome)):
        row_cells = tabela.add_row().cells
        row_cells[0].text = Nome[i]
        row_cells[1].text = DataDoAtendimento[i]
    """ 
    for mes, vendas, receita in dados_vendas:
        row_cells = tabela.add_row().cells
        row_cells[0].text = mes
        row_cells[1].text = vendas
        row_cells[2].text = receita
    """        
    

    # Gráficos ou Imagens (Exemplo, você pode adicionar imagens de gráficos se necessário)
    
    # documento.add_picture('caminho/para/seu/grafico.png')

    # Conclusão
    documento.add_heading('3. Conclusão', level=2)
    documento.add_paragraph(
        'Os resultados de vendas do período de Janeiro 2024 a Março 2024 mostram um crescimento contínuo, '
        'com um aumento significativo em Fevereiro.'
    )

    # Salvando o documento
    documento.save(f'Relatorio {hora_atual}.docx')
